"""
Document processor for chunking documents into smaller pieces.
"""

from typing import List, Dict, Any
from pathlib import Path
import re


class DocumentProcessor:
    """Process documents by chunking them into smaller pieces for embedding."""
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        supported_extensions: List[str] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_extensions = supported_extensions or [".txt", ".md", ".docx", ".pdf"]
    
    def _create_document(self, content: str, filepath: str) -> Any:
        """Create a document object."""
        # Simple document wrapper
        class SimpleDocument:
            def __init__(self, text, metadata):
                self.page_content = text
                self.metadata = metadata
        
        return SimpleDocument(
            text=content,
            metadata={"source": filepath, "filename": Path(filepath).name}
        )
    
    def read_docx(self, filepath: Path) -> str:
        """Extract text from DOCX file."""
        try:
            import docx
            doc = docx.Document(filepath)
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except ImportError:
            return ""
        except Exception as e:
            return ""
    
    def read_pdf(self, filepath: Path) -> str:
        """Extract text from PDF file."""
        try:
            import pypdf
            reader = pypdf.PdfReader(filepath)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
            return text
        except ImportError:
            return ""
        except Exception as e:
            return ""
    
    def read_txt(self, filepath: Path) -> str:
        """Read text file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            return ""
    
    def read_file(self, filepath: Path) -> str:
        """Read file based on extension."""
        ext = filepath.suffix.lower()
        
        if ext == '.docx':
            return self.read_docx(filepath)
        elif ext == '.pdf':
            return self.read_pdf(filepath)
        elif ext in ['.txt', '.md']:
            return self.read_txt(filepath)
        else:
            return ""
    
    def chunk_document(self, document) -> List[Any]:
        """Split document into chunks."""
        text = document.page_content
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # If paragraph alone is too long, split it further
            if len(para) > self.chunk_size:
                if current_chunk:
                    chunks.append(self._create_document(current_chunk, document.metadata.get("source", "")))
                    current_chunk = ""
                
                # Split long paragraph into sentences or fixed size
                sentences = re.split(r'(?<=[.!?])\s+', para)
                temp_chunk = ""
                for sentence in sentences:
                    if len(temp_chunk) + len(sentence) <= self.chunk_size:
                        temp_chunk += sentence + " "
                    else:
                        if temp_chunk:
                            chunks.append(self._create_document(temp_chunk.strip(), document.metadata.get("source", "")))
                        temp_chunk = sentence + " "
                if temp_chunk:
                    current_chunk = temp_chunk
            
            # Normal paragraph fits or current chunk plus paragraph fits
            elif len(current_chunk) + len(para) <= self.chunk_size:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
            else:
                # Current chunk is full, save it and start new
                if current_chunk:
                    chunks.append(self._create_document(current_chunk, document.metadata.get("source", "")))
                current_chunk = para
        
        # Add last chunk
        if current_chunk:
            chunks.append(self._create_document(current_chunk, document.metadata.get("source", "")))
        
        return chunks
    
    async def load_documents_from_directory(self, directory: str) -> List[Any]:
        """Load all documents from directory."""
        from pathlib import Path
        
        documents = []
        data_dir = Path(directory)
        
        for ext in self.supported_extensions:
            for filepath in data_dir.glob(f"*{ext}"):
                content = self.read_file(filepath)
                if content:
                    doc = self._create_document(content, str(filepath))
                    doc.name = filepath.name
                    documents.append(doc)
        
        return documents